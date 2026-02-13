import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Корневая папка проекта
ROOT_DIR = os.getcwd()

# Игнорируемые папки для дерева
IGNORED_DIRS = ['__pycache__', 'venv', 'migrations']

# Игнорируемые расширения/файлы для дерева
IGNORED_FILES = ['.pyc', '.pyo', '.db', '.sqlite3', '.pycache']

# Список ключевых файлов (относительные пути от корня проекта)
KEY_FILES = [
    'silant/settings.py',
    'silant/urls.py',
    'core/models.py',
    'core/views.py',
    'core/urls.py',
    'core/admin.py',
    'core/forms.py',
    'core/adapters.py',
    'core/templates/core/base.html',
    'core/templates/core/home.html',
    'core/templates/core/dashboard.html',
    'core/templates/core/machine_detail.html',
    'core/templates/core/machine_form.html',
]

# Выходной файл
OUTPUT_FILE = 'key_project_structure.docx'


def add_heading(doc, text, level=1):
    heading = doc.add_paragraph(text)
    heading.style.font.size = Pt(14 if level == 1 else 12)
    heading.style.font.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.style.paragraph_format.space_after = Pt(6)


def build_tree(path, prefix=''):
    tree = ''
    items = sorted(os.listdir(path))
    for i, item in enumerate(items):
        full_path = os.path.join(path, item)
        if any(ignored in item for ignored in IGNORED_DIRS + IGNORED_FILES):
            continue
        connector = '└── ' if i == len(items) - 1 else '├── '
        tree += f"{prefix}{connector}{item}\n"
        if os.path.isdir(full_path):
            extension = '    ' if i == len(items) - 1 else '│   '
            tree += build_tree(full_path, prefix + extension)
    return tree


def export_to_docx():
    doc = Document()

    # Раздел 1: Структура папок (дерево)
    add_heading(doc, "Структура проекта (дерево папок)", level=1)
    tree_str = build_tree(ROOT_DIR)
    add_code_block(doc, tree_str)

    # Раздел 2: Содержимое ключевых файлов
    add_heading(doc, "Содержимое ключевых файлов", level=1)

    for rel_path in KEY_FILES:
        full_path = os.path.join(ROOT_DIR, rel_path)

        if not os.path.exists(full_path):
            add_heading(doc, f"Файл: {rel_path} (не найден)", level=2)
            add_code_block(doc, "[Файл не найден или ещё не создан]")
            continue

        add_heading(doc, f"Файл: {rel_path}", level=2)

        try:
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
            add_code_block(doc, content)
        except UnicodeDecodeError:
            add_code_block(doc, "[Бинарный файл или не UTF-8 — содержимое не отображено]")
        except Exception as e:
            add_code_block(doc, f"[Ошибка чтения: {str(e)}]")

    doc.save(OUTPUT_FILE)
    print(f"Экспорт завершён! Файл сохранён как: {OUTPUT_FILE}")


if __name__ == '__main__':
    export_to_docx()